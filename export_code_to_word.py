import os
import sys
import tempfile
import shutil
import requests
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from pygments import highlight
from pygments.lexers import guess_lexer_for_filename
from pygments.formatters import HtmlFormatter
from bs4 import BeautifulSoup

def download_repo_zip(repo_url):
    if 'github.com' in repo_url or 'gitlab.com' in repo_url:
        if repo_url.endswith('/'):
            repo_url = repo_url[:-1]
        if '/tree/' in repo_url:
            repo_url = repo_url.replace('/tree/', '/archive/refs/heads/') + '.zip'
        elif '/blob/' in repo_url:
            repo_url = repo_url.split('/blob/')[0] + '/archive/refs/heads/main.zip'
        else:
            repo_url += '/archive/refs/heads/main.zip'
    else:
        raise ValueError('Only GitHub and GitLab are supported')

    response = requests.get(repo_url)
    if response.status_code != 200:
        raise ValueError(f'Failed to download repo: {response.status_code}')

    temp_dir = tempfile.mkdtemp()
    zip_path = os.path.join(temp_dir, 'repo.zip')
    with open(zip_path, 'wb') as f:
        f.write(response.content)

    shutil.unpack_archive(zip_path, temp_dir)
    os.remove(zip_path)

    # Найти первую распакованную папку
    subdirs = [os.path.join(temp_dir, d) for d in os.listdir(temp_dir)]
    repo_path = next((d for d in subdirs if os.path.isdir(d)), temp_dir)
    return repo_path

def add_code_to_docx(document, file_path, code):
    document.add_heading(file_path, level=2)
    # подсветка синтаксиса
    lexer = guess_lexer_for_filename(file_path, code)
    formatter = HtmlFormatter()
    highlighted_code = highlight(code, lexer, formatter)
    soup = BeautifulSoup(highlighted_code, 'html.parser')

    for line in soup.find_all('div', class_='highlight'):
        for pre in line.find_all('pre'):
            p = document.add_paragraph()
            run = p.add_run(pre.get_text())
            font = run.font
            font.name = 'Courier New'
            font.size = Pt(10)

def create_word_from_code(source_path, output_docx):
    document = Document()
    document.add_heading('Source Code Export', 0)

    for root, _, files in os.walk(source_path):
        for file in files:
            if file.endswith(('.py', '.cpp', '.c', '.java', '.js', '.ts', '.html', '.css', '.php')):
                full_path = os.path.join(root, file)
                with open(full_path, 'r', encoding='utf-8', errors='ignore') as f:
                    code = f.read()
                add_code_to_docx(document, os.path.relpath(full_path, source_path), code)

    document.save(output_docx)

if __name__ == "__main__":
    if len(sys.argv) != 3:
        print('Usage: python export_code_to_word.py <path_or_repo_url> <output.docx>')
        sys.exit(1)

    path_or_url = sys.argv[1]
    output_file = sys.argv[2]

    if path_or_url.startswith('http'):
        source_dir = download_repo_zip(path_or_url)
    else:
        source_dir = path_or_url

    create_word_from_code(source_dir, output_file)
    print(f"Document created: {output_file}")
